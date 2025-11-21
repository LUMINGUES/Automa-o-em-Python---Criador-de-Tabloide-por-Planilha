import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
# Importações necessárias para manipular o XML de travamento
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
import io
import os
import shutil
import glob

# --- Importação para PDF ---
# É necessário 'pip install docx2pdf'
try:
    from docx2pdf import convert as docx_to_pdf
except ImportError:
    docx_to_pdf = None # Define como None se a biblioteca não estiver instalada


# --- Configurações ---
PLANILHA_DADOS = 'dados.xlsx'
TEMPLATE_DOCX = 'template.docx'
SAIDA_PASTA = 'Documentos_Gerados'

# VOLTANDO AO LOTE: 16 produtos por documento/página
PRODUTOS_POR_DOC = 16

# O documento deve ter 4 colunas e 4 linhas (4x4 = 16 produtos)
CELULAS_POR_LINHA = 4
CELULAS_POR_COLUNA = 4

CARACTERES_MAXIMOS = 85 # Limite de caracteres para o TEXTPRODUCT (com ...)

# Altura e Largura da imagem em polegadas (aproximadamente 150x100 pixels)
IMG_WIDTH = Inches(1.50)
IMG_HEIGHT = Inches(1.0)

# Estilos de Fonte
FONT_NAME = 'Tahoma'
TEXT_SIZE = Pt(10)
PRICE_SIZE = Pt(12)
PRICE_COLOR = RGBColor(0xFF, 0x00, 0x00) # Vermelho (RGB)

# ------------------------------
# --- FUNÇÕES AUXILIARES ---
# ------------------------------

def limitar_texto(texto: str, max_chars: int) -> str:
    """Trunca o texto em 'max_chars' e adiciona '...' se for mais longo."""
    texto_str = str(texto)
    if len(texto_str) > max_chars:
        return texto_str[:max_chars] + '...'
    return texto_str

def formatar_preco_br(valor: Any) -> str:
    """Formata um valor para o padrão de moeda brasileiro R$ 0.000,00."""
    try:
        if isinstance(valor, str):
            valor_limpo = valor.replace('R$', '').replace(' ', '').replace('.', '').replace(',', '.')
            valor_float = float(valor_limpo)
        elif isinstance(valor, (int, float)):
            valor_float = float(valor)
        else:
            return str(valor)

        return f'R$ {valor_float:,.2f}'.replace(',', 'TEMP').replace('.', ',').replace('TEMP', '.')
    except:
        return str(valor)

def apply_style_and_replace(celula: Any, old_text: str, new_text: str, size: Pt, color: RGBColor = None):
    """Localiza, substitui e APLICA ESTILOS FORÇADOS (Tahoma, negrito, tamanho, cor)."""
    
    for p in celula.paragraphs:
        if old_text in p.text:
            p.text = ""
            new_run = p.add_run(new_text)
            new_run.font.name = FONT_NAME
            new_run.font.size = size
            new_run.bold = True
            
            if color:
                new_run.font.color.rgb = color
            
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            return

def insert_image_in_cell_at_marker(celula: Any, marker: str, url_img: str):
    """
    Localiza o marcador [IMG], insere a imagem centralizada no local do marcador e remove o marcador.
    """
    
    for p in celula.paragraphs:
        if marker in p.text:
            
            for run in p.runs:
                if marker in run.text:
                    text_parts = run.text.split(marker, 1)
                    run.text = ""
                    
                    if text_parts[0]:
                        run.text = text_parts[0]
                    
                    # --- Inserção da Imagem ---
                    try:
                        response = requests.get(url_img, stream=True, timeout=10)
                        response.raise_for_status()
                        image_stream = io.BytesIO(response.content)

                        img_run = p.add_run()
                        img_run.add_picture(image_stream, width=IMG_WIDTH, height=IMG_HEIGHT)
                        
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print(f"  -> Imagem inserida e redimensionada.")

                        if len(text_parts) > 1 and text_parts[1]:
                            p.add_run(text_parts[1], style=run.style)
                        
                        return
                        
                    except Exception as e:
                        print(f"  -> ERRO ao processar imagem: {e}")
                        p.add_run(f"[ERRO NA IMAGEM]").bold = True
                        if len(text_parts) > 1 and text_parts[1]:
                            p.add_run(text_parts[1], style=run.style)
                        return

def set_table_layout_properties(tabela: Any):
    """Aplica propriedades de layout para controlar o tamanho e a quebra de página."""
    
    tbl = tabela._tbl # Elemento XML da tabela
    
    # 1. DEFINIR LARGURA DA TABELA (Trava de Largura)
    tblPr = tbl.tblPr
    if tblPr.xpath('./w:tblW') == []:
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000') # 100% da largura da página
        tblPr.append(tblW)
        
    # 2. PROPRIEDADE PARA NÃO DIVIDIR LINHAS NA PÁGINA
    for row in tabela.rows:
        trPr = row._tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        cantSplit.set(qn('w:val'), '1')
        
        for child in trPr.xpath('./w:cantSplit'):
            trPr.remove(child)
            
        trPr.append(cantSplit)


# ------------------------------
# --- LÓGICA PRINCIPAL ---
# ------------------------------

def main():
    # 0. Configuração da Pasta de Saída
    if not os.path.exists(SAIDA_PASTA):
        try:
            os.makedirs(SAIDA_PASTA)
        except OSError as e:
            print(f"ERRO: Não foi possível criar a pasta de saída '{SAIDA_PASTA}'.")
            return
            
    # Lista para rastrear os nomes dos arquivos PDF gerados
    # Não é usada neste script, mas deixada para referência
    # arquivos_pdf_gerados = [] 

    # 1. Leitura dos Dados
    try:
        df = pd.read_excel(PLANILHA_DADOS)
        # Mapeamento robusto das colunas
        df.columns = [c.lower().strip() for c in df.columns]
        col_map = {c: 'name' for c in df.columns if 'nome' in c or 'name' in c}
        col_map.update({c: 'img_url' for c in df.columns if 'url' in c or 'img' in c})
        col_map.update({c: 'price' for c in df.columns if 'price' in c or 'preco' in c})
        df = df.rename(columns=col_map)
        df = df[['name', 'img_url', 'price']]
    except FileNotFoundError:
        print(f"ERRO: Arquivo de planilha '{PLANILHA_DADOS}' não encontrado.")
        return
    except Exception as e:
        print(f"ERRO ao ler a planilha: {e}")
        return

    total_produtos = len(df)
    if total_produtos == 0:
        print("Nenhum dado encontrado na planilha. Automação encerrada.")
        return
        
    num_docs = (total_produtos + PRODUTOS_POR_DOC - 1) // PRODUTOS_POR_DOC
    
    print(f"Início da automação: {total_produtos} produtos serão processados em {num_docs} documento(s) DOCX de 16 itens.")
    
    for i in range(num_docs):
        start_index = i * PRODUTOS_POR_DOC
        end_index = min((i + 1) * PRODUTOS_POR_DOC, total_produtos)
        
        produtos_do_lote = df.iloc[start_index:end_index].to_dict('records')
        
        # 2. Configurar nomes de arquivos
        nome_doc_docx = f'Catalogo_Lote_{i+1}.docx'
        nome_doc_pdf = f'Catalogo_Lote_{i+1}.pdf'
        caminho_doc_docx = os.path.join(SAIDA_PASTA, nome_doc_docx)
        caminho_doc_pdf = os.path.join(SAIDA_PASTA, nome_doc_pdf)
        
        try:
            # 2.1. Copiar o Template e carregar o documento
            print(f"\n[Lote {i+1}] Criando '{nome_doc_docx}'")
            shutil.copyfile(TEMPLATE_DOCX, caminho_doc_docx)
            doc = Document(caminho_doc_docx)
        except FileNotFoundError:
            print(f"ERRO CRÍTICO: Arquivo template '{TEMPLATE_DOCX}' não encontrado.")
            return
        except Exception as e:
            print(f"ERRO CRÍTICO ao iniciar o documento: {e}")
            return
            
        if not doc.tables:
            print(f"ERRO CRÍTICO: O documento '{TEMPLATE_DOCX}' não contém tabelas.")
            return

        tabela = doc.tables[0]
        # Aplica propriedades de layout na tabela
        set_table_layout_properties(tabela) 

        # 3. Processar os produtos do lote
        for j, produto in enumerate(produtos_do_lote):
            
            # Mapeamento da Célula Única (4x4)
            celula_col = j % CELULAS_POR_LINHA
            celula_row = j // CELULAS_POR_LINHA
            
            try:
                celula = tabela.cell(celula_row, celula_col)
            except IndexError:
                print(f"AVISO: Produto {j+1} fora dos limites da tabela. Ignorando.")
                break
                
            print(f" -> Lote {i+1}, Pos. {j+1}: {produto.get('name', 'N/A')}")
            
            # 4. Substituição de Texto (COM ESTILOS FORÇADOS)
            nome_limitado = limitar_texto(produto.get('name', ''), CARACTERES_MAXIMOS)
            apply_style_and_replace(celula, '[TEXTPRODUCT]', nome_limitado, TEXT_SIZE)
            
            preco_formatado = formatar_preco_br(produto.get('price', ''))
            apply_style_and_replace(celula, 'R$[PRICE]', preco_formatado, PRICE_SIZE, PRICE_COLOR)
            
            # 5. Inserção da Imagem
            url_img = produto.get('img_url', '')
            if url_img:
                insert_image_in_cell_at_marker(celula, '[IMG]', url_img)
            else:
                apply_style_and_replace(celula, '[IMG]', '[SEM IMAGEM]', TEXT_SIZE, PRICE_COLOR)
        
        # 6. Salvar o documento DOCX
        doc.save(caminho_doc_docx)
        print(f"Documento DOCX '{nome_doc_docx}' criado com sucesso.")

        # 7. Salvar o documento em PDF (Conversão individual)
        if docx_to_pdf:
            print(f"Convertendo DOCX para PDF: '{nome_doc_pdf}'")
            try:
                docx_to_pdf(caminho_doc_docx, caminho_doc_pdf)
                # arquivos_pdf_gerados.append(caminho_doc_pdf) # Não precisa mais desta linha
            except Exception as e:
                print(f"AVISO: Falha na conversão para PDF do {nome_doc_pdf}. Erro: {e}")
        else:
            print("AVISO: Módulo 'docx2pdf' não encontrado. Conversão para PDF ignorada.")

    print("\n--- VERIFICAÇÃO HUMANA NECESSÁRIA - FIM DA GERAÇÃO DE DOCUMENTOS. POR FAVOR VERIFIQUE ARQUIVO POR ARQUIVO E AO MODIFICAR ALGUM SALVE-O EM PDF DEPOIS Execute 'juntarpdf.py' para o ARQUIVO final. ---")
    


if __name__ == '__main__':
    main()